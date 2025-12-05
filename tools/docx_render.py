"""DOCX template rendering tool for Toolhub.

This module expects templates in a dedicated directory and renders new
DOCX files by replacing placeholders of the form ``{{KEY}}`` using values
provided in the payload. It is designed for use via the Toolhub /run
endpoint so other services (e.g. n8n) can trigger document generation.
"""

import os
import time
from datetime import datetime
from typing import Dict, Iterable

from docx import Document
from loguru import logger


TEMPLATES_DIR = os.getenv("DOCX_TEMPLATES_DIR", "/data/templates")
OUTPUT_DIR = os.getenv("DOCX_OUTPUT_DIR", "/data/output")
LOG_PATH = os.getenv("DOCX_RENDER_LOG", "/logs/docx-render.log")
LOG_LEVEL = os.getenv("DOCX_RENDER_LOG_LEVEL", "DEBUG")

_logger_configured = False


def _configure_logger() -> None:
    """Ensure loguru writes to the dedicated log file once."""
    global _logger_configured
    if _logger_configured:
        return

    os.makedirs(os.path.dirname(LOG_PATH), exist_ok=True)
    logger.remove()
    logger.add(LOG_PATH, level=LOG_LEVEL, rotation=None, retention=None)
    _logger_configured = True


def _safe_filename(name: str, field: str) -> str:
    """Reject absolute paths or traversal attempts in filenames."""
    if not name or not isinstance(name, str):
        raise ValueError(f"{field} must be a non-empty string")

    if os.path.isabs(name):
        raise ValueError(f"{field} must be a filename, not an absolute path")

    if any(sep in name for sep in (os.sep, os.altsep) if sep):
        raise ValueError(f"{field} must not contain path separators")

    normalized = os.path.normpath(name)
    if normalized.startswith("..") or ".." in normalized.split(os.sep):
        raise ValueError(f"{field} must not contain parent directory traversal")

    return name


def replace_placeholders(doc: Document, data: Dict[str, str]) -> Dict[str, int]:
    """Replace ``{{KEY}}`` placeholders in all paragraphs and table cells.

    Returns a count of replacements performed per key for debug logging.
    """

    replacement_counts: Dict[str, int] = {key: 0 for key in data}

    def _replace_text(text: str, replacements: Iterable[tuple[str, str]]) -> tuple[str, Dict[str, int]]:
        local_counts: Dict[str, int] = {}
        for key, value in replacements:
            placeholder = f"{{{{{key}}}}}"
            occurrences = text.count(placeholder)
            if occurrences:
                text = text.replace(placeholder, value)
                replacement_counts[key] = replacement_counts.get(key, 0) + occurrences
                local_counts[key] = local_counts.get(key, 0) + occurrences
        return text, local_counts

    replacements = data.items()

    for paragraph in doc.paragraphs:
        original = paragraph.text
        updated, local_counts = _replace_text(original, replacements)
        if updated != original:
            logger.debug("Updated paragraph with placeholders", extra={"counts": local_counts})
            paragraph.text = updated

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                original = cell.text
                updated, local_counts = _replace_text(original, replacements)
                if updated != original:
                    logger.debug(
                        "Updated table cell with placeholders",
                        extra={
                            "counts": local_counts,
                            "table": table_index,
                            "row": row_index,
                            "cell": cell_index,
                        },
                    )
                    cell.text = updated

    return replacement_counts


def handler(payload: Dict) -> Dict:
    """Entry point for the Toolhub docx-render tool."""
    _configure_logger()

    logger.info("docx-render invoked")
    try:
        template_name = payload.get("template") if isinstance(payload, dict) else None
        output_name = payload.get("output_name") if isinstance(payload, dict) else None
        data = payload.get("data") if isinstance(payload, dict) else None

        if not template_name:
            raise ValueError("'template' is required and must be a filename")
        if not isinstance(data, dict):
            raise ValueError("'data' must be a dictionary of placeholder values")

        safe_template = _safe_filename(str(template_name), "template")
        safe_output = _safe_filename(str(output_name), "output_name") if output_name else None

        logger.debug(
            "Validated payload for docx-render",
            extra={
                "template": safe_template,
                "output_name": safe_output,
                "data_keys": sorted([str(key) for key in data.keys()]),
            },
        )

        template_path = os.path.join(TEMPLATES_DIR, safe_template)
        if not os.path.isfile(template_path):
            logger.error(f"Template not found at path: {template_path}")
            return {
                "status": "error",
                "error": {
                    "type": "FileNotFoundError",
                    "message": f"Template not found: {safe_template}",
                },
            }

        if not safe_output:
            stem, _ = os.path.splitext(safe_template)
            timestamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
            safe_output = f"{stem}_out_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, safe_output)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        logger.info("Rendering DOCX template")
        logger.info(f"Template path: {template_path}")
        logger.info(f"Output path: {output_path}")

        document = Document(template_path)
        replacement_counts = replace_placeholders(document, {k: str(v) for k, v in data.items()})
        logger.debug(
            "Placeholder replacement summary",
            extra={"replacements": {k: v for k, v in replacement_counts.items() if v}},
        )
        document.save(output_path)

        logger.info(f"Rendered DOCX written to {output_path}")
        return {
            "status": "ok",
            "template": safe_template,
            "output_path": output_path,
            "placeholders_filled": sorted([str(key) for key in data.keys()]),
            "timestamp": int(time.time()),
        }
    except Exception as exc:  # noqa: BLE001
        logger.error(f"docx-render error: {exc.__class__.__name__}: {exc}")
        return {
            "status": "error",
            "error": {
                "type": exc.__class__.__name__,
                "message": str(exc),
            },
        }
