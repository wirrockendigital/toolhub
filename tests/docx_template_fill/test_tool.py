from __future__ import annotations

import importlib
import os
import tempfile
import unittest
from pathlib import Path

from docx import Document

import mcp_tools.docx_template_fill.renderer as renderer_module
import mcp_tools.docx_template_fill.tool as tool_module


class DocxTemplateFillTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tempdir = tempfile.TemporaryDirectory()
        self.template_root = Path(self.tempdir.name) / "templates"
        self.output_root = Path(self.tempdir.name) / "output"
        self.template_root.mkdir(parents=True, exist_ok=True)
        self.output_root.mkdir(parents=True, exist_ok=True)

        os.environ["DOCX_TEMPLATE_ROOT"] = str(self.template_root)
        os.environ["DOCX_OUTPUT_ROOT"] = str(self.output_root)

        # Reload modules to pick up new environment-based roots.
        self.renderer = importlib.reload(renderer_module)
        self.tool = importlib.reload(tool_module)

    def tearDown(self) -> None:
        self.tempdir.cleanup()

    def _write_template(self, name: str, paragraphs: list[str]) -> Path:
        path = self.template_root / name
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    def _read_paragraphs(self, path: Path) -> list[str]:
        doc = Document(path)
        return [p.text for p in doc.paragraphs]

    def test_valid_generation(self) -> None:
        self._write_template("example.docx", ["Hello {{NAME}}", "Week {{WEEK}} report."])

        payload = {
            "template": "example.docx",
            "data": {"NAME": "Alice", "WEEK": "42"},
            "output_subdir": "reports/2025",
            "output_filename": "summary.docx",
        }

        result = self.tool.fill_docx_template(payload)

        output_path = Path(result["output_file"])
        self.assertTrue(output_path.exists())
        self.assertIn("Hello Alice", " ".join(self._read_paragraphs(output_path)))
        self.assertIn("Week 42", " ".join(self._read_paragraphs(output_path)))

    def test_missing_template_raises(self) -> None:
        payload = {
            "template": "missing.docx",
            "data": {"NAME": "Bob"},
            "output_filename": "out.docx",
        }
        with self.assertRaises(self.renderer.TemplateValidationError):
            self.tool.fill_docx_template(payload)

    def test_invalid_template_extension(self) -> None:
        payload = {
            "template": "template.txt",
            "data": {"NAME": "Bob"},
            "output_filename": "out.docx",
        }
        with self.assertRaises(self.renderer.TemplateValidationError):
            self.tool.fill_docx_template(payload)

    def test_invalid_output_filename(self) -> None:
        self._write_template("example.docx", ["Hello {{NAME}}"])
        payload = {
            "template": "example.docx",
            "data": {"NAME": "Bob"},
            "output_filename": "bad/name.docx",
        }
        with self.assertRaises(self.renderer.TemplateValidationError):
            self.tool.fill_docx_template(payload)

    def test_invalid_output_subdir(self) -> None:
        self._write_template("example.docx", ["Hi"])
        payload = {
            "template": "example.docx",
            "data": {"NAME": "Bob"},
            "output_filename": "out.docx",
            "output_subdir": "../escape",
        }
        with self.assertRaises(self.renderer.TemplateValidationError):
            self.tool.fill_docx_template(payload)

    def test_creates_nested_output_dir(self) -> None:
        self._write_template("example.docx", ["Hi {{NAME}}"])
        payload = {
            "template": "example.docx",
            "data": {"NAME": "Bob"},
            "output_filename": "out.docx",
            "output_subdir": "nested/dir",
        }
        result = self.tool.fill_docx_template(payload)
        output_path = Path(result["output_file"])
        self.assertTrue(output_path.exists())
        self.assertTrue(output_path.parent.exists())

    def test_overwrite_blocked(self) -> None:
        template_path = self._write_template("example.docx", ["Hello"])
        payload = {
            "template": template_path.name,
            "data": {"NAME": "Bob"},
            "output_filename": "out.docx",
        }
        # First render succeeds
        first_output = self.tool.fill_docx_template(payload)
        self.assertTrue(Path(first_output["output_file"]).exists())
        # Second render with same path should fail
        with self.assertRaises(self.renderer.TemplateValidationError):
            self.tool.fill_docx_template(payload)

    def test_placeholder_mismatch_warns_only(self) -> None:
        self._write_template("example.docx", ["Hello {{NAME}}", "Role {{ROLE}}"])
        payload = {
            "template": "example.docx",
            "data": {"NAME": "Bob"},
            "output_filename": "out.docx",
        }
        result = self.tool.fill_docx_template(payload)
        self.assertIn("ROLE", result["placeholders_missing"])
        self.assertTrue(Path(result["output_file"]).exists())


if __name__ == "__main__":  # pragma: no cover
    unittest.main()
